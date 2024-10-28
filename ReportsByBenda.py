import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
import docx.shared
import datetime
import requests  # For fetching the logo from a URL

# Function to fetch the logo from a URL
def fetch_image_from_url(url):
    response = requests.get(url)
    if response.status_code == 200:
        return BytesIO(response.content)
    else:
        st.error("Failed to fetch the logo from the provided URL.")
        return None

# Function to generate the Word document with graphs and a logo
def generate_report_with_logo(data, client_name, logo_url):
    doc = Document()

    # Ensure date is in datetime format and extract the month
    data['Date OF Purchase'] = pd.to_datetime(data['Date OF Purchase'], format="%d/%m/%Y")
    data['Month'] = data['Date OF Purchase'].dt.to_period('M')

    # Clean and convert 'Profit Per Sale' to float, handling errors gracefully
    data['Profit Per Sale Cleaned'] = pd.to_numeric(
        data['Profit Per Sale'].str.replace('[\$,]', '', regex=True),
        errors='coerce'
    ).fillna(0)

    # Calculate total metrics
    total_sales = data['eBay Price'].sum()
    total_profit_per_sale = data['Profit Per Sale Cleaned'].sum()

    # Fetch the logo from the URL
    logo_image = fetch_image_from_url(logo_url)

    # Add the logo to the header of every page
    if logo_image:
        section = doc.sections[0]
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = 1  # Center alignment
        run = header_paragraph.add_run()
        run.add_picture(logo_image, width=docx.shared.Inches(2.0))

    # Create and append graphs
    fig_list = []

    # Plot 1: Number of Sales per Month
    monthly_sales_count = data.groupby('Month').size()
    fig1, ax1 = plt.subplots()
    ax1.bar(monthly_sales_count.index.astype(str), monthly_sales_count, color="purple")
    ax1.set_title("Number of Sales per Month")
    ax1.set_xlabel("Month")
    ax1.set_ylabel("Number of Sales")
    ax1.tick_params(axis='x', rotation=45)
    fig_list.append(fig1)

    # Plot 2: Sum of Sales per Month
    monthly_sales_sum = data.groupby('Month')['eBay Price'].sum()
    fig2, ax2 = plt.subplots()
    ax2.bar(monthly_sales_sum.index.astype(str), monthly_sales_sum, color="blue")
    ax2.set_title("Sum of Sales per Month")
    ax2.set_xlabel("Month")
    ax2.set_ylabel("Sum of Sales (eBay Price)")
    ax2.tick_params(axis='x', rotation=45)
    fig_list.append(fig2)

    # Plot 3: Sum of Profit per Month
    monthly_profit_sum = data.groupby('Month')['Profit Per Sale Cleaned'].sum()
    fig3, ax3 = plt.subplots()
    ax3.bar(monthly_profit_sum.index.astype(str), monthly_profit_sum, color="green")
    ax3.set_title("Sum of Profit per Month")
    ax3.set_xlabel("Month")
    ax3.set_ylabel("Sum of Profit")
    ax3.tick_params(axis='x', rotation=45)
    fig_list.append(fig3)

    # Plot 4: Profit Per Sale Over Time
    fig4, ax4 = plt.subplots()
    ax4.plot(data['Date OF Purchase'], data['Profit Per Sale Cleaned'], color="orange")
    ax4.set_title("Profit Per Sale Over Time")
    ax4.set_xlabel("Date of Purchase")
    ax4.set_ylabel("Profit Per Sale ($)")
    ax4.tick_params(axis='x', rotation=45)
    fig_list.append(fig4)

    # Plot 5: eBay Price Over Time
    fig5, ax5 = plt.subplots()
    ax5.plot(data['Date OF Purchase'], data['eBay Price'], color="blue")
    ax5.set_title("eBay Price Over Time")
    ax5.set_xlabel("Date of Purchase")
    ax5.set_ylabel("eBay Price ($)")
    ax5.tick_params(axis='x', rotation=45)
    fig_list.append(fig5)

    # Plot 6: Average Profit per Product per Month
    avg_profit_per_product = monthly_profit_sum / monthly_sales_count
    fig6, ax6 = plt.subplots()
    ax6.plot(avg_profit_per_product.index.astype(str), avg_profit_per_product, color="red")
    ax6.set_title("Average Profit per Product per Month")
    ax6.set_xlabel("Month")
    ax6.set_ylabel("Average Profit")
    ax6.tick_params(axis='x', rotation=45)
    fig_list.append(fig6)

    # Add summary to the Word document
    doc.add_heading(f'Investment Report for {client_name}', 0)
    doc.add_paragraph(f"Total Sales: ${total_sales:.2f}")
    doc.add_paragraph(f"Total Profit (Per Sale): ${total_profit_per_sale:.2f}")

    # Save each graph into the Word document
    for i, fig in enumerate(fig_list):
        img_stream = BytesIO()
        fig.savefig(img_stream, format='png')
        img_stream.seek(0)
        
        # Add the graph to the document
        doc.add_paragraph(f'Graph {i+1}:')
        doc.add_picture(img_stream)

    # Save the document to a buffer and return it
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# Streamlit app setup
st.title('Investment Report Generator')

st.write("Upload a CSV file with your sales data and get a Word document with analysis and graphs, including your logo.")

uploaded_file = st.file_uploader("Choose a CSV file", type="csv")

if uploaded_file:
    # Read the CSV data
    data = pd.read_csv(uploaded_file)
    
    # Input for client's name
    client_name = st.text_input("Enter the client's name", value="Client")
    
    # Input for the logo URL
    logo_url = st.text_input("Enter the URL of your logo", value="https://i.postimg.cc/kgvhv1Mn/BENDA-logo-black.png")
    
    if st.button("Generate Report"):
        # Generate the report with the provided data and logo URL
        report_buffer = generate_report_with_logo(data, client_name, logo_url)
        
        # Provide a download link for the report
        st.download_button(
            label="Download Report",
            data=report_buffer,
            file_name=f"{client_name}_Investment_Report_{datetime.date.today()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
